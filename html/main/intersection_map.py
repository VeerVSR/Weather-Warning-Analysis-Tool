import geopandas as gpd
import pandas as pd
import matplotlib.pyplot as plt
import os
import numpy as np
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import matplotlib.image as mpimg
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import matplotlib.colors as mcolors
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.gridspec as gridspec
import sys
import json
import psycopg2

# PostgreSQL connection settings
db_config = {
    'dbname': 'imd',
    'user': 'postgres',
    'password': 'sql123T',
    'host': 'localhost',
    'port': 5432
}

# Connect to PostgreSQL
conn = psycopg2.connect(**db_config)
cur = conn.cursor()
cur.execute("""
    SELECT geojson_data FROM imd_geojson
    ORDER BY submitted_at DESC
    LIMIT 1
""")
row = cur.fetchone()
if not row:
    print("❌ No GeoJSON found in the database.")
    exit(1)
data = row[0]
cur.close()
conn.close()

output_dir = r'F:\nginx-1.28.0\html\imd\processed'

warning_icons = {
    "Heavy Rain": "icons/004-rain-1.png",
    "Very Heavy Rain": "icons/003-rain.png",
    "Extremely Heavy Rain": "icons/002-rainy.png",
    "Heavy Snow": "icons/005-snow.png",
    "Thunderstorm & Lightning": "icons/006-thunder.png",
    "Hailstorm": "icons/007-hailstorm.png",
    "Dust Storm": "icons/008-dust-storm.png",
    "Dust Raising Winds": "icons/009-dust-raising-winds.png",
    "Strong Surface Winds": "icons/010-strong-surface-winds.png",
    "Heat Wave": "icons/011-heat-wave.png",
    "Hot Day": "icons/014-hot-day.png",
    "Hot and Humid": "icons/012-humidity.png",
    "Warm Night": "icons/013-warm-night.png",
    "Cold Wave": "icons/003-rain.png",
    "Cold Day": "icons/003-rain.png",
    "Ground Frost": "icons/015-ice.png",
    "Fog": "icons/016-fog.png",
}

def color_to_hex(color_name):
    try:
        rgb = mcolors.to_rgb(color_name)
        return mcolors.to_hex(rgb)
    except ValueError:
        return '#FF0000'

def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def filter_intersections(multipolygon_gdf, other_gdf, threshold=0.1):
    results = []
    for idx, multipolygon in multipolygon_gdf.iterrows():
        for jdx, other_polygon in other_gdf.iterrows():
            intersection = multipolygon.geometry.intersection(other_polygon.geometry)
            if not intersection.is_empty:
                intersection_area = intersection.area
                original_area = multipolygon.geometry.area
                if (intersection_area / original_area) > threshold:
                    result = multipolygon.copy()
                    result['intersection_area'] = intersection_area
                    result['intersection_ratio'] = intersection_area / original_area
                    result['Warning'] = other_polygon['Warning']
                    result['Date'] = other_polygon['Date']
                    result['Color'] = color_to_hex(other_polygon.get('Color', 'red'))
                    results.append(result)
    return gpd.GeoDataFrame(results, columns=multipolygon_gdf.columns.tolist() + ['intersection_area', 'intersection_ratio', 'Warning', 'Date', 'Color'])

def add_district_labels(ax, gdf, label_col='District'):
    for _, row in gdf.iterrows():
        if row.geometry.is_empty:
            continue
        try:
            centroid = row.geometry.centroid
            ax.text(centroid.x, centroid.y, str(row[label_col]), fontsize=6, ha='center', va='center', color='white',
                    bbox=dict(facecolor='black', alpha=0.5, boxstyle='round,pad=0.2'))
        except Exception as e:
            print(f"Labeling error: {e}")

def plot_warnings(multipolygon_gdf, filtered_gdf, dates, output_dir, base_name):
    for state in ['PUNJAB', 'HARYANA']:
        fig = plt.figure(figsize=(15, 22))
        gs = gridspec.GridSpec(4, 2, height_ratios=[0.1, 1, 1, 1])

        title_ax = fig.add_subplot(gs[0, :])
        title_ax.axis('off')
        title_ax.text(0.5, 0.5, f"District wise weather warnings for {state}\nDate {dates[0].strftime('%d %B %Y')}",
                      ha='center', va='center', fontsize=16, fontweight='bold')

        for i, date in enumerate(dates):
            ax = fig.add_subplot(gs[i // 2 + 1, i % 2])
            state_gdf = multipolygon_gdf[multipolygon_gdf['STATE'] == state]
            state_gdf.plot(ax=ax, edgecolor='black', facecolor='darkgreen')

            warning_areas = filtered_gdf[(filtered_gdf['STATE'] == state) & (filtered_gdf['Date'] == date)]

            for _, area in warning_areas.iterrows():
                area_color = area['Color']
                if area.geometry.geom_type == 'Polygon':
                    ax.fill(area.geometry.exterior.xy[0], area.geometry.exterior.xy[1],
                            facecolor=area_color, edgecolor='black', alpha=0.5)
                elif area.geometry.geom_type == 'MultiPolygon':
                    for polygon in area.geometry.geoms:
                        ax.fill(polygon.exterior.xy[0], polygon.exterior.xy[1],
                                facecolor=area_color, edgecolor='black', alpha=0.5)

            for _, area in warning_areas.iterrows():
                centroid = area.geometry.centroid
                warning = area['Warning']
                if warning and warning in warning_icons:
                    icon_path = warning_icons[warning]
                    if os.path.exists(icon_path):
                        icon = mpimg.imread(icon_path)
                        imagebox = OffsetImage(icon, zoom=0.3)
                        ab = AnnotationBbox(imagebox, (centroid.x, centroid.y + 0.1), 
                                            xybox=(0, 0), boxcoords="offset points", frameon=False)
                        ax.add_artist(ab)

            add_district_labels(ax, state_gdf)
            ax.set_title(f"Day {i + 1}\n{date.strftime('%d %B %Y')}")
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.tick_params(labelbottom=False, labelleft=False)

        legend_ax = fig.add_subplot(gs[3, 1])
        legend_ax.axis('off')
        legend_items = [
            plt.Rectangle((0, 0), 1, 1, fc="darkgreen", ec="black", label="No Warning"),
            plt.Rectangle((0, 0), 1, 1, fc="yellow", ec="black", label="Be Updated"),
            plt.Rectangle((0, 0), 1, 1, fc="orange", ec="black", label="Be Prepared"),
            plt.Rectangle((0, 0), 1, 1, fc="red", ec="black", label="Take Action")
        ]
        legend_ax.legend(handles=legend_items, loc='center', title="Warning Level", fontsize=10, title_fontsize=12)

        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, f"{base_name}_{state}_warnings.png"), dpi=300, bbox_inches='tight')
        plt.close(fig)

def main(input_file, warning_file, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_file = os.path.join(output_dir, f"{base_name}_filtered.geojson")

    multipolygon_gdf = gpd.read_file(input_file)
    other_gdf = gpd.read_file(warning_file)

    if multipolygon_gdf.crs != other_gdf.crs:
        other_gdf = other_gdf.to_crs(multipolygon_gdf.crs)

    filtered_gdf = filter_intersections(multipolygon_gdf, other_gdf)

    if not filtered_gdf.empty and filtered_gdf.crs:
        filtered_gdf.crs = filtered_gdf.crs.to_wkt()
    filtered_gdf.to_file(output_file, driver='GeoJSON')

    if not filtered_gdf.empty:
        district_col = 'District' if 'District' in filtered_gdf.columns else 'DISTRICT_L'
        doc = Document()
        doc.add_heading('District Warnings Table', 0)
        current_date = pd.to_datetime(filtered_gdf['Date'].min()).to_pydatetime()

        all_districts = multipolygon_gdf.copy()
        all_districts['District'] = all_districts['District'].str.upper()
        filtered_gdf['District'] = filtered_gdf['District'].str.upper()

        for state in ['PUNJAB', 'HARYANA']:
            doc.add_heading(state, level=1)
            state_districts = all_districts[all_districts['STATE'] == state].copy()
            state_filtered = filtered_gdf[filtered_gdf['STATE'] == state].copy()

            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            header_cells = table.rows[0].cells
            header_cells[0].text = 'क्षेत्र/जिले\nArea/Districts'
            for i in range(1, 6):
                date = current_date + timedelta(days=i - 1)
                header_cells[i].text = f'{date.strftime("%d-%m-%y")}\nचेतावनी/Warning'

            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True

            for _, district_row in state_districts.iterrows():
                district_name = district_row['District']
                row_cells = table.add_row().cells
                row_cells[0].text = district_name

                for i in range(1, 6):
                    date = current_date + timedelta(days=i - 1)
                    match = state_filtered[
                        (state_filtered['District'] == district_name) &
                        (pd.to_datetime(state_filtered['Date']).dt.date == date.date())
                    ]
                    if not match.empty:
                        row_cells[i].text = match['Warning'].values[0]
                        set_cell_background(row_cells[i], match['Color'].values[0].lstrip('#'))
                    else:
                        row_cells[i].text = 'NIL'
                        set_cell_background(row_cells[i], '00FF00')

                    row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)

            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(1.5)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)

        word_file = os.path.join(output_dir, f"{base_name}_warnings_table.docx")
        doc.save(word_file)

        # Convert DOCX to PDF
        import subprocess
        pdf_file = word_file.replace('.docx', '.pdf')
        subprocess.run([
            r'C:\Program Files\LibreOffice\program\soffice.exe', '--headless',
            '--convert-to', 'pdf', '--outdir', output_dir, word_file
        ], check=True)

        print(f"Warnings table converted to PDF: {pdf_file}")

        dates = [current_date + timedelta(days=i) for i in range(5)]
        plot_warnings(multipolygon_gdf, filtered_gdf, dates, output_dir, base_name)
    else:
        print("No intersections found with area greater than the threshold.")

if __name__ == "__main__":
    output_dir = r'F:\nginx-1.28.0\html\imd\processed'
    os.makedirs(output_dir, exist_ok=True)
    warning_file = os.path.join(output_dir, "latest_warning.geojson")
    with open(warning_file, 'w') as f:
        json.dump(data, f)
    input_file = r'F:\nginx-1.28.0\html\imd\merged_district_punjab_haryana.geojson'
    main(input_file, warning_file, output_dir)